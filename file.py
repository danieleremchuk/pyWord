# file.py - file management, manipulation and validation
import docx
import csv
from datetime import datetime

def do_Doc(my_doc):
    doc = docx.Document(my_doc)
    numPar = 0
    for _ in doc.paragraphs:
        numPar = numPar+1
    return numPar

def do_CSV(my_par, my_file, my_brew):
    name = my_brew
    now = datetime.now()

    month = now.month
    day = now.day
    year = now.year
    hour = now.hour
    minute = now.minute
    second = now.second

    dt = str(month)+str(day)+str(year)+"_"+str(hour)+str(minute)+str(second)
    filename = "bpr_"+name.replace(" ", "")+"_"+dt.replace(" ", "")+".csv"
    csv_file = open(filename, "w")
    my_par = do_Doc(my_file)
    doc = docx.Document(my_file)

    csv_file.write("iParNum,vParText")
    csv_file.write("\n")

    j=1
    while j < my_par:
        csv_file.write(str(j))
        csv_file.write(",")
        csv_file.write(doc.paragraphs[j-1].text)
        if j < my_par-1:
            csv_file.write("\n")
        j = j+1

    return print("all done!")

def valid_file(my_file):
    if len(my_file) == 0:
        return False
    if ".doc" not in my_file:
        return False
    else:
        return True

def valid_name(my_name):
    if len(my_name) == 0:
        return False
    else:
        return True
